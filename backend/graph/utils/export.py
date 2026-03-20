"""
export.py

Exports the final proposal letter as a formatted Word document (.docx).

Usage (run from backend/):
    python -m graph.utils.export <letter.txt> -o proposal_letter.docx

Or programmatically:
    from graph.utils.export import letter_to_docx
    docx_bytes = letter_to_docx(letter_text)
"""

import argparse
import io
import os
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ENV_LOC = "../.env"


def letter_to_docx(letter_text: str, firm_name: str = None) -> bytes:
    """Convert a proposal transmittal letter to a formatted Word document.

    Args:
        letter_text: The letter as plain text
        firm_name: Optional firm name for the header

    Returns:
        The .docx file as bytes
    """
    doc = Document()

    # Page margins — 1 inch all around
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Parse the letter into lines
    lines = letter_text.strip().split("\n")

    for line in lines:
        stripped = line.strip()

        if not stripped:
            # Empty line — add spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(14)
            continue

        # Check if this is the closing signature block
        # Lines like "Sincerely," or "Respectfully submitted," get extra space before
        if stripped.lower().startswith(("sincerely", "respectfully", "regards", "best regards", "very truly yours")):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(stripped)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)

        # Apply font to the run
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)

    # Convert to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def letter_to_docx_file(letter_text: str, output_path: str, firm_name: str = None):
    """Convert a letter to a .docx file and save it.

    Args:
        letter_text: The letter as plain text
        output_path: Where to save the .docx file
        firm_name: Optional firm name for the header
    """
    docx_bytes = letter_to_docx(letter_text, firm_name)
    with open(output_path, "wb") as f:
        f.write(docx_bytes)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Export a proposal letter to a Word document."
    )
    parser.add_argument(
        "letter_file",
        help="Path to a text file containing the letter",
    )
    parser.add_argument(
        "--output", "-o",
        default="proposal_letter.docx",
        help="Output .docx file path (default: proposal_letter.docx)",
    )

    args = parser.parse_args()

    with open(args.letter_file) as f:
        letter_text = f.read()

    letter_to_docx_file(letter_text, args.output)
    print(f"Saved to: {args.output}")


if __name__ == "__main__":
    main()
